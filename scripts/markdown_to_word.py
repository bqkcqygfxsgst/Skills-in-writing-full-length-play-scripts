#!/usr/bin/env python3
"""
剧本Word格式生成器 — Markdown → .docx
遵循标准话剧剧本格式
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class DramaToWord:
    def __init__(self, md_path, output_path):
        self.md_path = Path(md_path)
        self.output_path = Path(output_path)
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """设置Word文档样式"""
        # 设置正常文本样式
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)
    
    def read_markdown(self):
        """读取Markdown文件"""
        return self.md_path.read_text(encoding='utf-8')
    
    def add_title(self, text):
        """添加标题"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.name = '黑体'
    
    def add_subtitle(self, text):
        """添加副标题"""
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(11)
    
    def add_scene_header(self, text):
        """添加场景头"""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(11)
    
    def add_character_dialogue(self, character, dialogue):
        """添加角色对白"""
        # 角色名
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        r = p.add_run(character + ':')
        r.font.bold = True
        r.font.size = Pt(10.5)
        
        # 对白内容
        for line in dialogue.split('\n'):
            if line.strip():
                p = self.doc.add_paragraph(line)
                p.paragraph_format.left_indent = Inches(1.0)
                p.paragraph_format.first_line_indent = Inches(0.5)
                for run in p.runs:
                    run.font.size = Pt(10.5)
    
    def add_stage_direction(self, text):
        """添加舞台指示"""
        text = text.strip('() []')
        p = self.doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(1.0)
        for run in p.runs:
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
    
    def add_normal_text(self, text):
        """添加普通文本"""
        if text.strip():
            p = self.doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(10.5)
    
    def convert(self):
        """转换Markdown为Word"""
        md_content = self.read_markdown()
        
        # 提取标题
        title_match = re.search(r'^#\s+(.+?)(?:\n|$)', md_content)
        if title_match:
            self.add_title(title_match.group(1))
        
        # 分行处理
        lines = md_content.split('\n')
        i = 0
        current_character = None
        
        while i < len(lines):
            line = lines[i]
            
            # 跳过YAML frontmatter和Markdown元数据
            if line.startswith('---') or line.startswith('```'):
                i += 1
                continue
            
            # 标题处理 (## 开头)
            if line.startswith('## '):
                text = line.replace('## ', '').strip()
                self.add_scene_header(text)
            
            # 场景头处理 [时间：...]
            elif line.startswith('[时间：') or line.startswith('[地点：'):
                self.add_scene_header(line)
            
            # 舞台指示 (...)
            elif line.startswith('(') or (line.strip().startswith('(') and line.strip().endswith(')')):
                self.add_stage_direction(line.strip())
            
            # 角色对白
            elif ':' in line and not line.startswith('**') and re.match(r'^[^\s：][^\：]*：', line):
                # 分离角色和对白
                colon_pos = line.find('：') if '：' in line else line.find(':')
                if colon_pos > 0:
                    character = line[:colon_pos].strip()
                    dialogue = line[colon_pos+1:].strip()
                    
                    # 收集多行对白
                    full_dialogue = dialogue
                    i += 1
                    while i < len(lines):
                        next_line = lines[i]
                        # 如果下一行不是新的角色或场景标记，就是延续对白
                        if (not next_line.startswith('[') and 
                            not next_line.startswith('##') and
                            not re.match(r'^[^\s：][^\：]*：', next_line) and
                            next_line.strip()):
                            if next_line.startswith('  ('):
                                # 这是舞台指示
                                break
                            full_dialogue += '\n' + next_line
                            i += 1
                        else:
                            break
                    
                    self.add_character_dialogue(character, full_dialogue)
                    i -= 1  # 因为后面会++，所以这里--
            
            # 普通文本
            elif line.strip() and not line.startswith('#'):
                self.add_normal_text(line)
            
            i += 1
        
        # 保存文档
        self.doc.save(str(self.output_path))
        print(f"✅ Word文档已生成: {self.output_path}")


def main():
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python markdown_to_word.py <Markdown文件>")
        sys.exit(1)
    
    md_path = sys.argv[1]
    output_path = md_path.replace('.md', '.docx')
    
    converter = DramaToWord(md_path, output_path)
    converter.convert()
    
    print("\n" + "="*60)
    print("📄 剧本Word格式转换完成!")
    print("="*60)


if __name__ == '__main__':
    main()
